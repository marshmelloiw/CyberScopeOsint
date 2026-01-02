#!/usr/bin/env python3
"""
Script to generate Test Plan Report for CyberScope OSINT Platform
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    return heading

def add_table_with_headers(doc, headers, rows):
    """Create a table with headers and rows"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def main():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Test Plan Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('CyberScope OSINT Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Senior Graduation Project')
    doc.add_paragraph('Computer Engineering Department')
    doc.add_paragraph('')
    
    # 1. Introduction
    add_heading_with_style(doc, '1. Introduction', 1)
    
    doc.add_paragraph(
        'This test plan document outlines the testing strategy and approach for CyberScope, '
        'an Open Source Intelligence (OSINT) platform designed to help security teams detect '
        'threats early by consolidating data from multiple security tools and providing unified '
        'visibility, alerts, and AI-powered reporting.'
    )
    
    add_heading_with_style(doc, '1.1 Purpose of the Test Plan', 2)
    doc.add_paragraph(
        'The main purpose of this test plan is to ensure that the CyberScope platform works '
        'correctly and reliably before it is deployed. Testing helps us find bugs, verify that '
        'all features work as expected, and make sure the system is secure and performs well '
        'under normal usage conditions.'
    )
    doc.add_paragraph(
        'This document serves as a guide for the testing team (which in this case is mainly me, '
        'the developer) to systematically test all parts of the application. It helps organize '
        'the testing process and makes sure nothing important gets missed.'
    )
    
    add_heading_with_style(doc, '1.2 Brief Project Description', 2)
    doc.add_paragraph(
        'CyberScope is a web-based OSINT platform that integrates with multiple security tools '
        'and services to provide comprehensive threat intelligence. The system consists of two '
        'main parts: a React-based frontend and a Spring Boot backend.'
    )
    doc.add_paragraph(
        'The frontend is built with React 18, Vite, and Tailwind CSS. It provides a modern, '
        'responsive user interface where users can create scans, view reports, manage API keys, '
        'and monitor security threats. The backend is a Java Spring Boot application that handles '
        'authentication, scan execution, integration with external OSINT services, and report generation.'
    )
    doc.add_paragraph(
        'Key features include user authentication with JWT tokens and multi-factor authentication (MFA), '
        'scan management for different types of security checks (email, domain, IP address analysis), '
        'integration with services like Shodan, VirusTotal, Have I Been Pwned, and Twitter API, '
        'automated report generation using AI (Gemini), and a notification system for alerting users '
        'about important events.'
    )
    
    add_heading_with_style(doc, '1.3 Why Testing is Needed', 2)
    doc.add_paragraph(
        'Testing is crucial for this project because CyberScope deals with security-sensitive data '
        'and integrates with multiple external services. If the system has bugs or security vulnerabilities, '
        'it could lead to incorrect threat assessments, data leaks, or unauthorized access.'
    )
    doc.add_paragraph(
        'Since this is a graduation project that will be evaluated by professors and potentially '
        'demonstrated to others, it needs to work reliably. Testing helps ensure that all the features '
        'I implemented actually work as intended and that the user experience is smooth and intuitive.'
    )
    doc.add_paragraph(
        'Additionally, testing helps catch integration issues early. The platform connects to many '
        'external APIs (Shodan, VirusTotal, etc.), and these integrations can fail in various ways. '
        'Testing helps identify problems with API calls, error handling, and data processing before '
        'they become bigger issues.'
    )
    
    add_heading_with_style(doc, '1.4 Testing Goals', 2)
    doc.add_paragraph(
        'The primary goals of testing for this project are:'
    )
    
    goals = [
        'Verify that all user authentication features work correctly, including login, registration, password reset, and MFA setup',
        'Ensure that scan creation and execution work properly for different scan types (email, domain, IP)',
        'Validate that integrations with external OSINT services (Shodan, VirusTotal, HIBP, Twitter) function as expected',
        'Check that the user interface is responsive and works well on different screen sizes',
        'Verify that role-based access control (Admin, Analyst, Viewer) is properly enforced',
        'Test that report generation (PDF and HTML formats) produces correct and readable output',
        'Ensure that the notification system correctly alerts users about scan completions and other events',
        'Validate that API key management works securely and prevents unauthorized access',
        'Check that error handling works properly when things go wrong (network failures, invalid inputs, etc.)',
        'Verify basic security measures like input validation, SQL injection prevention, and XSS protection'
    ]
    
    for i, goal in enumerate(goals, 1):
        p = doc.add_paragraph(goal, style='List Bullet')
    
    # 2. Scope
    add_heading_with_style(doc, '2. Scope', 1)
    
    add_heading_with_style(doc, '2.1 What Will Be Tested', 2)
    doc.add_paragraph(
        'The following components and features will be thoroughly tested:'
    )
    
    tested_items = [
        ('Authentication Module', 'Login, registration, password reset, MFA setup and verification, JWT token handling'),
        ('Dashboard', 'Statistics display, recent scans, notifications summary, navigation'),
        ('Scan Management', 'Creating new scans, viewing scan history, scan detail pages, scan status updates'),
        ('OSINT Integrations', 'Shodan IP lookups, VirusTotal domain/IP analysis, HIBP email breach checks, Twitter API integration'),
        ('Report Generation', 'PDF report generation, HTML report generation, report viewing and download'),
        ('User Management', 'User CRUD operations (Admin only), role assignment, user status management'),
        ('API Key Management', 'API key creation, viewing, updating, and deletion (Admin only)'),
        ('Notifications', 'Notification display, marking as read, notification preferences'),
        ('Settings', 'User profile updates, password changes, MFA configuration'),
        ('UI/UX Components', 'Responsive design, form validation, error messages, loading states')
    ]
    
    for item, desc in tested_items:
        p = doc.add_paragraph(f'{item}: ', style='List Bullet')
        p.add_run(desc).italic = True
    
    add_heading_with_style(doc, '2.2 What Will Not Be Tested', 2)
    doc.add_paragraph(
        'Due to time constraints and project scope, the following items are out of scope for this test plan:'
    )
    
    not_tested = [
        'Load testing with thousands of concurrent users (basic performance testing will be done, but not extensive stress testing)',
        'Penetration testing by external security experts (basic security checks will be performed, but not a full security audit)',
        'Cross-browser testing on all possible browsers (testing will focus on Chrome, Firefox, and Edge)',
        'Mobile app testing (the responsive web interface will be tested, but no native mobile apps exist)',
        'Third-party service reliability (we will test our integration code, but cannot test the external services themselves)',
        'Database migration testing for all possible migration scenarios (basic migration testing will be done)',
        'Accessibility compliance testing beyond basic checks (WCAG full compliance is not in scope)'
    ]
    
    for item in not_tested:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_with_style(doc, '2.3 Limitations', 2)
    doc.add_paragraph(
        'Several limitations affect the testing process:'
    )
    doc.add_paragraph(
        'First, testing is being done by a single person (the developer), which means it might take longer '
        'and some edge cases could be missed. Having multiple testers would be ideal, but for a graduation '
        'project, this is a reasonable limitation.'
    )
    doc.add_paragraph(
        'Second, some external services have rate limits or require paid API keys. This means we cannot '
        'test all possible scenarios exhaustively. For example, we might not be able to test what happens '
        'when we hit Shodan\'s rate limit because we want to avoid using up our API quota during testing.'
    )
    doc.add_paragraph(
        'Third, the testing environment might differ from a production environment. The database might be '
        'smaller, network conditions might be different, and we might not have the same infrastructure setup. '
        'This is normal for a student project, but it means some issues might only appear in production.'
    )
    doc.add_paragraph(
        'Finally, time is a constraint. As a graduation project, there is a deadline, so we need to balance '
        'thoroughness with practicality. We will focus on testing the most critical features and common use cases.'
    )
    
    add_heading_with_style(doc, '2.4 Assumptions', 2)
    doc.add_paragraph(
        'The following assumptions are made for this test plan:'
    )
    
    assumptions = [
        'The test environment will have internet connectivity to access external OSINT APIs',
        'All required API keys for external services (Shodan, VirusTotal, etc.) will be available and valid',
        'The PostgreSQL database will be set up and accessible',
        'Test users with different roles (Admin, Analyst, Viewer) will be created in the database',
        'The development team (me) has basic knowledge of testing tools and techniques',
        'The application will be tested on Windows 10/11, which is the primary development environment',
        'Modern web browsers (Chrome, Firefox, Edge) will be used for frontend testing',
        'Postman or similar tools will be available for API testing'
    ]
    
    for assumption in assumptions:
        doc.add_paragraph(assumption, style='List Bullet')
    
    # 3. Test Items
    add_heading_with_style(doc, '3. Test Items', 1)
    
    add_heading_with_style(doc, '3.1 Modules and Components', 2)
    doc.add_paragraph(
        'The CyberScope platform consists of several main modules that need to be tested:'
    )
    
    modules = [
        ('Authentication Module', 
         'Handles user login, registration, password reset, and MFA. Components: Login page, Register page, '
         'Forgot Password page, Reset Password page, MFA Setup component, MFA Verification component. '
         'Input: Email, password, MFA codes. Output: JWT tokens, user session, authentication status.'),
        
        ('Dashboard Module',
         'Displays overview statistics and recent activity. Components: Dashboard page, StatWidget components, '
         'recent scans list, notifications summary. Input: User authentication token. Output: Statistics data, '
         'recent scans, notification counts.'),
        
        ('Scan Management Module',
         'Manages security scans for emails, domains, and IP addresses. Components: ScansList page, NewScan page, '
         'ScanDetail page. Input: Scan type, target (email/domain/IP), scan parameters. Output: Scan results, '
         'scan status, scan history.'),
        
        ('Report Generation Module',
         'Generates PDF and HTML reports from scan results using AI analysis. Components: Reports page, '
         'GeminiReportDetail page. Input: Scan ID, report format preference. Output: PDF or HTML report files.'),
        
        ('User Management Module',
         'Allows admins to manage users. Components: UserManagement page. Input: User data, role assignments. '
         'Output: User list, updated user information.'),
        
        ('API Key Management Module',
         'Manages API keys for external services. Components: APIKeys page. Input: API key data, service name. '
         'Output: API key list, encrypted key storage.'),
        
        ('Notification Module',
         'Displays and manages user notifications. Components: Notifications page. Input: User ID. '
         'Output: Notification list, read/unread status.'),
        
        ('Settings Module',
         'User profile and preference management. Components: Settings page. Input: Profile updates, '
         'password changes, MFA settings. Output: Updated user profile, confirmation messages.')
    ]
    
    for module, desc in modules:
        p = doc.add_paragraph(f'{module}: ', style='List Bullet')
        p.add_run(desc)
    
    add_heading_with_style(doc, '3.2 Key Features', 2)
    doc.add_paragraph(
        'The following key features represent the core functionality of the platform:'
    )
    
    features = [
        'User authentication with JWT tokens and optional MFA',
        'Role-based access control (Admin, Analyst, Viewer)',
        'Multi-type security scanning (email, domain, IP address)',
        'Integration with Shodan for IP and service discovery',
        'Integration with VirusTotal for domain and IP analysis',
        'Integration with Have I Been Pwned for email breach checking',
        'Integration with Twitter API for social media intelligence',
        'AI-powered report generation using Google Gemini',
        'PDF and HTML report export',
        'Real-time scan status updates',
        'Notification system for scan completions and alerts',
        'Responsive web interface that works on mobile, tablet, and desktop',
        'API key management for external service integrations',
        'User management for administrators',
        'Dashboard with statistics and overview'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_with_style(doc, '3.3 Input-Output Descriptions', 2)
    doc.add_paragraph(
        'Understanding the inputs and outputs of each module helps in designing effective test cases:'
    )
    
    io_descriptions = [
        ('Login', 
         'Input: Email (string), Password (string). Output: JWT access token, refresh token, user role, user ID.'),
        ('Create Scan',
         'Input: Scan type (email/domain/IP), target value (string), optional parameters. Output: Scan ID, scan status (started).'),
        ('Get Scan Results',
         'Input: Scan ID (UUID). Output: Scan status, results data (JSON), risk score, completion timestamp.'),
        ('Generate Report',
         'Input: Scan ID, report format (PDF/HTML). Output: Report file (binary) or report URL.'),
        ('Update User',
         'Input: User ID, updated fields (name, role, status). Output: Updated user object, success message.'),
        ('Create API Key',
         'Input: Service name, API key value, description. Output: Encrypted API key ID, creation timestamp.'),
        ('Get Notifications',
         'Input: User ID, optional filters. Output: List of notification objects with read/unread status.')
    ]
    
    for io_name, io_desc in io_descriptions:
        p = doc.add_paragraph(f'{io_name}: ', style='List Bullet')
        p.add_run(io_desc).italic = True
    
    # 4. Test Strategy
    add_heading_with_style(doc, '4. Test Strategy', 1)
    
    add_heading_with_style(doc, '4.1 Functional Testing', 2)
    doc.add_paragraph(
        'Functional testing is the most important type of testing for this project. It involves checking '
        'that each feature works the way it is supposed to work. For example, when a user logs in with '
        'correct credentials, they should be redirected to the dashboard. When they enter wrong credentials, '
        'they should see an error message.'
    )
    doc.add_paragraph(
        'I will test all the main user flows: registration, login, creating a scan, viewing results, '
        'generating reports, and managing users (for admins). Each feature will be tested with both valid '
        'inputs (happy path) and invalid inputs (error cases). For instance, I will test creating a scan '
        'with a valid email address, and also test what happens when someone tries to create a scan with '
        'an invalid email format or an empty field.'
    )
    
    add_heading_with_style(doc, '4.2 UI/UX Testing', 2)
    doc.add_paragraph(
        'UI/UX testing focuses on how the application looks and feels to users. This includes checking '
        'that buttons are in logical places, forms are easy to fill out, error messages are clear and helpful, '
        'and the overall design is consistent throughout the application.'
    )
    doc.add_paragraph(
        'I will test the responsive design by resizing the browser window and checking that the layout adapts '
        'properly on mobile, tablet, and desktop sizes. I will also verify that all interactive elements '
        '(buttons, links, forms) work correctly and provide appropriate feedback. For example, when a user '
        'clicks a button to start a scan, the button should show a loading state so the user knows something '
        'is happening. Navigation should be intuitive, and users should always know where they are in the application.'
    )
    
    add_heading_with_style(doc, '4.3 Performance Testing (Basic)', 2)
    doc.add_paragraph(
        'Performance testing checks how fast the application responds and whether it can handle normal usage '
        'without becoming slow or unresponsive. For a graduation project, I will do basic performance testing '
        'rather than extensive load testing.'
    )
    doc.add_paragraph(
        'I will measure how long it takes for pages to load, how quickly API calls complete, and whether the '
        'application remains responsive when performing multiple operations. For example, I will check that the '
        'dashboard loads within a reasonable time (under 2-3 seconds), and that starting a scan doesn\'t freeze '
        'the interface. I will also test with a moderate number of scans in the database (maybe 50-100) to see '
        'if the scan list page still loads quickly. If something takes too long, I will investigate and optimize if possible.'
    )
    
    add_heading_with_style(doc, '4.4 Security Testing (Simple Checks)', 2)
    doc.add_paragraph(
        'Security testing is important because CyberScope deals with security data and user authentication. '
        'I will perform basic security checks to ensure common vulnerabilities are not present.'
    )
    doc.add_paragraph(
        'I will test for SQL injection by trying to input SQL code in form fields and seeing if it gets executed. '
        'I will check for XSS (Cross-Site Scripting) by trying to inject JavaScript code and verifying it doesn\'t '
        'run. I will verify that authentication is required for protected routes and that users cannot access '
        'admin features without the proper role. I will also check that passwords are properly hashed (not stored '
        'in plain text) and that JWT tokens expire correctly. These are basic checks that a student can perform '
        'without specialized security tools.'
    )
    
    add_heading_with_style(doc, '4.5 Regression Testing', 2)
    doc.add_paragraph(
        'Regression testing means checking that existing features still work after making changes or adding new features. '
        'This is important because sometimes fixing one bug can break something else.'
    )
    doc.add_paragraph(
        'I will maintain a list of critical user flows (like login, creating a scan, viewing reports) and test these '
        'flows after each major change. For example, if I add a new feature to the dashboard, I will make sure that '
        'login still works, scans can still be created, and reports can still be generated. This doesn\'t have to be '
        'exhaustive every time, but the main features should be verified to prevent breaking changes.'
    )
    
    add_heading_with_style(doc, '4.6 Integration Testing', 2)
    doc.add_paragraph(
        'Integration testing checks that different parts of the system work together correctly. In CyberScope, this '
        'mainly involves testing the integration between the frontend and backend, and between the backend and external APIs.'
    )
    doc.add_paragraph(
        'I will test that API calls from the frontend correctly reach the backend and return the expected data. I will '
        'verify that when the frontend requests scan results, the backend properly queries the database and returns the '
        'data in the correct format. I will also test the integrations with external services: when a scan is created, '
        'does the backend correctly call Shodan or VirusTotal? Does it handle API errors gracefully? Does it store the '
        'results in the database correctly? These tests ensure the whole system works as a cohesive unit.'
    )
    
    # 5. Test Environment
    add_heading_with_style(doc, '5. Test Environment', 1)
    
    add_heading_with_style(doc, '5.1 Hardware and Software', 2)
    doc.add_paragraph(
        'The testing will be performed on the following hardware and software setup:'
    )
    
    env_items = [
        ('Operating System', 'Windows 10/11 (development machine)'),
        ('CPU', 'Modern multi-core processor (sufficient for running local servers)'),
        ('RAM', 'At least 8GB (to run frontend, backend, and database simultaneously)'),
        ('Storage', 'Sufficient space for Node.js, Java, PostgreSQL, and project files'),
        ('Network', 'Internet connection for accessing external APIs and downloading dependencies'),
        ('Java Version', 'Java 17 or higher (required for Spring Boot)'),
        ('Node.js Version', 'Node.js 18 or higher (required for React frontend)'),
        ('Database', 'PostgreSQL 12 or higher'),
        ('Build Tools', 'Maven (for Java backend), npm (for Node.js frontend)')
    ]
    
    for item, desc in env_items:
        p = doc.add_paragraph(f'{item}: ', style='List Bullet')
        p.add_run(desc)
    
    add_heading_with_style(doc, '5.2 Testing Tools', 2)
    doc.add_paragraph(
        'The following tools will be used for testing:'
    )
    
    tools = [
        ('Web Browser', 'Google Chrome (primary), Mozilla Firefox, Microsoft Edge for cross-browser testing'),
        ('Postman', 'For testing REST API endpoints directly, checking request/response formats, and testing error cases'),
        ('Browser DevTools', 'For inspecting network requests, checking console errors, and debugging frontend issues'),
        ('PostgreSQL Client', 'pgAdmin or command-line psql for verifying database state and checking data integrity'),
        ('Git', 'For version control and tracking changes during testing'),
        ('IDE', 'IntelliJ IDEA or VS Code for code inspection and debugging')
    ]
    
    for tool, desc in tools:
        p = doc.add_paragraph(f'{tool}: ', style='List Bullet')
        p.add_run(desc)
    
    add_heading_with_style(doc, '5.3 Test Accounts', 2)
    doc.add_paragraph(
        'The following test accounts will be created with different roles to test role-based access control:'
    )
    
    test_accounts = [
        ('admin@test.com', 'Admin123!', 'ADMIN', 'Full system access, can manage users and API keys'),
        ('analyst@test.com', 'Analyst123!', 'ANALYST', 'Can create scans and view reports, cannot manage users'),
        ('viewer@test.com', 'Viewer123!', 'VIEWER', 'Read-only access, can view reports but cannot create scans'),
        ('testuser@test.com', 'Test123!', 'ANALYST', 'General purpose test account for functional testing')
    ]
    
    headers = ['Email', 'Password', 'Role', 'Purpose']
    add_table_with_headers(doc, headers, test_accounts)
    
    add_heading_with_style(doc, '5.4 Setup Steps', 2)
    doc.add_paragraph(
        'Before testing can begin, the test environment needs to be set up. Here are the steps:'
    )
    
    setup_steps = [
        'Install PostgreSQL and create a database named "cyberscope"',
        'Run database migrations using Flyway (migrations are in backend/src/main/resources/db/migration)',
        'Configure backend application.yml with database connection details and API keys',
        'Start the Spring Boot backend server (runs on port 8080 by default)',
        'Install Node.js dependencies for the frontend (npm install in frontend directory)',
        'Configure frontend environment variables if needed',
        'Start the frontend development server (npm run dev, runs on port 5173)',
        'Verify both servers are running and accessible',
        'Create test user accounts in the database (can be done via registration or directly in database)',
        'Verify external API keys are configured and valid (Shodan, VirusTotal, etc.)'
    ]
    
    for i, step in enumerate(setup_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # 6. Test Case Design
    add_heading_with_style(doc, '6. Test Case Design', 1)
    
    add_heading_with_style(doc, '6.1 Test Case Preparation Approach', 2)
    doc.add_paragraph(
        'Test cases were prepared by analyzing the requirements and breaking down each feature into testable scenarios. '
        'I started by identifying the main user flows and then created test cases for both successful operations and error cases.'
    )
    doc.add_paragraph(
        'For each feature, I thought about: What should happen when everything works correctly? What should happen '
        'when the user provides invalid input? What should happen when something goes wrong (like a network error)? '
        'This approach helps ensure comprehensive coverage of each feature.'
    )
    
    add_heading_with_style(doc, '6.2 Test Design Techniques', 2)
    doc.add_paragraph(
        'I used several simple test design techniques:'
    )
    
    techniques = [
        ('Equivalence Partitioning', 
         'This means grouping similar inputs together. For example, all valid email addresses belong to one '
         'equivalence class, and all invalid email addresses belong to another. Instead of testing every possible '
         'email, I test one from each class.'),
        ('Boundary Value Analysis',
         'This involves testing values at the edges of valid ranges. For example, if a password must be at least '
         '8 characters, I test with 7 characters (invalid), 8 characters (valid boundary), and 9 characters (valid).'),
        ('Positive and Negative Testing',
         'Positive testing checks that valid inputs produce correct results. Negative testing checks that invalid '
         'inputs are properly rejected with appropriate error messages.'),
        ('User Story Based Testing',
         'I created test cases based on user stories like "As a user, I want to log in so I can access the dashboard." '
         'This ensures tests align with actual user needs.')
    ]
    
    for technique, desc in techniques:
        p = doc.add_paragraph(f'{technique}: ', style='List Bullet')
        p.add_run(desc)
    
    add_heading_with_style(doc, '6.3 Sample Test Cases', 2)
    doc.add_paragraph(
        'The following table contains sample test cases covering the main features of the application. '
        'These test cases will be executed during the testing phase.'
    )
    
    # Sample test cases
    test_cases = [
        ('TC-001', 'User Login with Valid Credentials',
         'User account exists in database, backend server is running',
         '1. Navigate to /auth/login\n2. Enter valid email and password\n3. Click Login button',
         'User is redirected to dashboard, JWT token is stored, user session is established'),
        
        ('TC-002', 'User Login with Invalid Credentials',
         'Backend server is running',
         '1. Navigate to /auth/login\n2. Enter invalid email or password\n3. Click Login button',
         'Error message displayed: "Invalid email or password", user remains on login page'),
        
        ('TC-003', 'User Registration with Valid Data',
         'Backend server is running, email is not already registered',
         '1. Navigate to /auth/register\n2. Fill in all required fields with valid data\n3. Click Register button',
         'User account is created, success message displayed, user can log in'),
        
        ('TC-004', 'User Registration with Invalid Email Format',
         'Backend server is running',
         '1. Navigate to /auth/register\n2. Enter invalid email format (e.g., "notanemail")\n3. Fill other fields\n4. Click Register',
         'Validation error displayed for email field, registration is prevented'),
        
        ('TC-005', 'Create Email Scan',
         'User is logged in as Analyst or Admin, backend is running',
         '1. Navigate to /dashboard/scans/new\n2. Select "Email" as scan type\n3. Enter valid email address\n4. Click Start Scan',
         'Scan is created, scan ID is returned, scan status shows as "started"'),
        
        ('TC-006', 'Create Domain Scan',
         'User is logged in as Analyst or Admin, Shodan API key is configured',
         '1. Navigate to /dashboard/scans/new\n2. Select "Domain" as scan type\n3. Enter valid domain (e.g., "example.com")\n4. Click Start Scan',
         'Scan is created, backend initiates Shodan lookup, scan status updates'),
        
        ('TC-007', 'View Scan Results',
         'A scan exists in database with completed status',
         '1. Navigate to /dashboard/scans\n2. Click on a completed scan\n3. View scan detail page',
         'Scan results are displayed, risk score is shown, OSINT data is visible'),
        
        ('TC-008', 'Generate PDF Report',
         'A completed scan exists, Gemini API key is configured',
         '1. Navigate to scan detail page\n2. Click "Generate Report" button\n3. Select PDF format',
         'PDF report is generated and downloaded, report contains scan results and AI analysis'),
        
        ('TC-009', 'Admin Views User List',
         'User is logged in as Admin',
         '1. Navigate to /dashboard/users\n2. View user management page',
         'List of all users is displayed with roles and status'),
        
        ('TC-010', 'Admin Creates New User',
         'User is logged in as Admin',
         '1. Navigate to /dashboard/users\n2. Click "Add User" button\n3. Fill in user details and select role\n4. Click Save',
         'New user is created, appears in user list, can log in with provided credentials'),
        
        ('TC-011', 'Viewer Cannot Create Scan',
         'User is logged in as Viewer role',
         '1. Navigate to /dashboard/scans/new',
         'User is redirected to /403 Forbidden page, error message displayed'),
        
        ('TC-012', 'Password Reset Request',
         'User account exists with valid email',
         '1. Navigate to /auth/forgot-password\n2. Enter registered email address\n3. Click Submit',
         'Password reset email is sent (if email service configured), success message displayed'),
        
        ('TC-013', 'MFA Setup',
         'User is logged in, MFA is not yet enabled',
         '1. Navigate to /dashboard/settings\n2. Click "Enable MFA" button\n3. Scan QR code with authenticator app\n4. Enter verification code',
         'MFA is enabled, user must provide MFA code on next login'),
        
        ('TC-014', 'View Dashboard Statistics',
         'User is logged in, some scans exist in database',
         '1. Navigate to /dashboard\n2. View dashboard page',
         'Statistics are displayed (total scans, completed scans, etc.), recent scans list is shown'),
        
        ('TC-015', 'API Key Management (Admin Only)',
         'User is logged in as Admin',
         '1. Navigate to /dashboard/apikeys\n2. Click "Add API Key"\n3. Enter service name and API key value\n4. Click Save',
         'API key is encrypted and stored, appears in API keys list'),
        
        ('TC-016', 'View Notifications',
         'User is logged in, some notifications exist',
         '1. Navigate to /dashboard/notifications\n2. View notifications page',
         'List of notifications is displayed, unread notifications are highlighted'),
        
        ('TC-017', 'Invalid Scan Type Handling',
         'User is logged in as Analyst',
         '1. Navigate to /dashboard/scans/new\n2. Try to submit form without selecting scan type\n3. Click Start Scan',
         'Validation error displayed, scan is not created'),
        
        ('TC-018', 'Network Error Handling',
         'Backend server is stopped or unreachable',
         '1. User is logged in\n2. Try to create a scan\n3. Observe error handling',
         'User-friendly error message displayed, application does not crash'),
        
        ('TC-019', 'Responsive Design - Mobile View',
         'Application is running, browser window can be resized',
         '1. Open application in browser\n2. Resize window to mobile size (375px width)\n3. Navigate through pages',
         'Layout adapts to mobile size, all features are accessible, no horizontal scrolling'),
        
        ('TC-020', 'Session Expiration Handling',
         'User is logged in, JWT token expires',
         '1. Wait for token to expire (or manually expire it)\n2. Try to access protected route',
         'User is redirected to login page, session is cleared')
    ]
    
    headers = ['Test ID', 'Test Description', 'Preconditions', 'Steps', 'Expected Result']
    add_table_with_headers(doc, headers, test_cases)
    
    # 7. Test Schedule
    add_heading_with_style(doc, '7. Test Schedule', 1)
    
    doc.add_paragraph(
        'The testing phase is planned to take approximately 3-4 weeks, running parallel with final development '
        'and bug fixes. The schedule is flexible to accommodate unexpected issues, but follows this general timeline:'
    )
    
    add_heading_with_style(doc, '7.1 Week 1: Setup and Core Functionality Testing', 2)
    doc.add_paragraph(
        'During the first week, the focus is on setting up the test environment and testing the most critical '
        'functionality: authentication and basic scan operations.'
    )
    
    week1_tasks = [
        'Day 1-2: Set up test environment, create test accounts, verify all services are running',
        'Day 3-4: Test authentication module (login, registration, password reset) - Test cases TC-001 to TC-004, TC-012',
        'Day 5: Test basic scan creation and viewing - Test cases TC-005, TC-007'
    ]
    
    for task in week1_tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    add_heading_with_style(doc, '7.2 Week 2: Feature Testing and Integration', 2)
    doc.add_paragraph(
        'The second week focuses on testing all major features and verifying integrations work correctly.'
    )
    
    week2_tasks = [
        'Day 1-2: Test scan management for all scan types (email, domain, IP) - Test cases TC-006, TC-017',
        'Day 3: Test OSINT service integrations (verify Shodan, VirusTotal, HIBP calls work)',
        'Day 4: Test report generation (PDF and HTML) - Test case TC-008',
        'Day 5: Test user management and API key management (Admin features) - Test cases TC-009, TC-010, TC-015'
    ]
    
    for task in week2_tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    add_heading_with_style(doc, '7.3 Week 3: UI/UX, Security, and Edge Cases', 2)
    doc.add_paragraph(
        'The third week covers UI/UX testing, basic security checks, and testing edge cases and error scenarios.'
    )
    
    week3_tasks = [
        'Day 1-2: UI/UX testing - responsive design, form validation, error messages - Test case TC-019',
        'Day 3: Security testing - SQL injection, XSS, authentication bypass attempts',
        'Day 4: Test role-based access control thoroughly - Test case TC-011',
        'Day 5: Test error handling and edge cases - Test cases TC-018, TC-020'
    ]
    
    for task in week3_tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    add_heading_with_style(doc, '7.4 Week 4: Regression Testing and Final Checks', 2)
    doc.add_paragraph(
        'The final week involves regression testing, performance checks, and preparing test deliverables.'
    )
    
    week4_tasks = [
        'Day 1-2: Regression testing - re-test all critical user flows after bug fixes',
        'Day 3: Basic performance testing - page load times, API response times',
        'Day 4: Final testing of notifications and settings - Test cases TC-013, TC-014, TC-016',
        'Day 5: Document test results, prepare bug reports, write test summary'
    ]
    
    for task in week4_tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    add_heading_with_style(doc, '7.5 Milestones', 2)
    doc.add_paragraph(
        'Key milestones in the testing schedule:'
    )
    
    milestones = [
        ('Milestone 1: Test Environment Ready', 'End of Week 1, Day 2 - All test accounts created, environment verified'),
        ('Milestone 2: Core Features Tested', 'End of Week 2 - Authentication and scan functionality verified'),
        ('Milestone 3: All Features Tested', 'End of Week 3 - All modules tested, security checks completed'),
        ('Milestone 4: Testing Complete', 'End of Week 4 - All test cases executed, bugs documented, test summary prepared')
    ]
    
    for milestone, desc in milestones:
        p = doc.add_paragraph(f'{milestone}: ', style='List Bullet')
        p.add_run(desc)
    
    # 8. Roles and Responsibilities
    add_heading_with_style(doc, '8. Roles and Responsibilities', 1)
    
    doc.add_paragraph(
        'Since this is a graduation project developed by a single student, the roles are simplified compared '
        'to a commercial project. However, it is still useful to define responsibilities:'
    )
    
    add_heading_with_style(doc, '8.1 Tester', 2)
    doc.add_paragraph(
        'The tester (which is also the developer in this case) is responsible for:'
    )
    
    tester_responsibilities = [
        'Executing test cases according to the test plan',
        'Documenting test results and any bugs found',
        'Reporting issues to the developer (self) with clear descriptions and steps to reproduce',
        'Retesting fixed bugs to verify they are resolved',
        'Maintaining test data and test accounts',
        'Preparing test deliverables (test summary, bug reports)'
    ]
    
    for resp in tester_responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    add_heading_with_style(doc, '8.2 Developer', 2)
    doc.add_paragraph(
        'The developer (also me) is responsible for:'
    )
    
    developer_responsibilities = [
        'Fixing bugs reported during testing',
        'Providing clarification on expected behavior when test results are unclear',
        'Ensuring the test environment is properly configured',
        'Updating documentation if requirements change',
        'Implementing any missing features needed for testing'
    ]
    
    for resp in developer_responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    add_heading_with_style(doc, '8.3 Supervisor', 2)
    doc.add_paragraph(
        'The project supervisor (professor) is responsible for:'
    )
    
    supervisor_responsibilities = [
        'Reviewing the test plan and providing feedback',
        'Evaluating the testing approach and test coverage',
        'Assessing the quality of test deliverables',
        'Providing guidance on testing best practices if needed'
    ]
    
    for resp in supervisor_responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    add_heading_with_style(doc, '8.4 Team Member Roles (If Applicable)', 2)
    doc.add_paragraph(
        'If other team members are involved (though this appears to be an individual project), they could help with:'
    )
    
    team_roles = [
        'Peer review of test cases',
        'Independent testing of specific modules',
        'User acceptance testing from an end-user perspective'
    ]
    
    for role in team_roles:
        doc.add_paragraph(role, style='List Bullet')
    
    # 9. Risks and Mitigations
    add_heading_with_style(doc, '9. Risks and Mitigations', 1)
    
    doc.add_paragraph(
        'Several risks could affect the testing process. This section identifies potential issues and how they will be addressed.'
    )
    
    add_heading_with_style(doc, '9.1 Technical Risks', 2)
    
    tech_risks = [
        ('External API Unavailability',
         'Risk: External services like Shodan or VirusTotal might be down or have rate limits that prevent testing.',
         'Mitigation: Use mock responses or test stubs when external services are unavailable. Document which tests require live APIs.'),
        
        ('Database Issues',
         'Risk: Database might become corrupted or migrations might fail, requiring time to fix.',
         'Mitigation: Keep database backups, use a separate test database, and have scripts to reset the database to a clean state.'),
        
        ('Environment Configuration Problems',
         'Risk: Test environment might not match development environment, causing issues that are hard to reproduce.',
         'Mitigation: Document all configuration steps clearly, use environment variables for configuration, and keep a setup checklist.'),
        
        ('Integration Failures',
         'Risk: Frontend and backend might have version mismatches or communication issues.',
         'Mitigation: Test API endpoints independently with Postman first, then test frontend integration. Keep API documentation updated.')
    ]
    
    for risk_title, risk_desc, mitigation in tech_risks:
        p = doc.add_paragraph(f'{risk_title}: ', style='List Bullet')
        p.add_run(risk_desc)
        p = doc.add_paragraph(f'Mitigation: {mitigation}', style='List Bullet 2')
    
    add_heading_with_style(doc, '9.2 Time and Communication Risks', 2)
    
    time_risks = [
        ('Testing Takes Longer Than Expected',
         'Risk: Finding and fixing bugs might take more time than planned, delaying the project.',
         'Mitigation: Prioritize critical bugs first, document non-critical issues for later, and be flexible with the schedule.'),
        
        ('Insufficient Test Coverage',
         'Risk: Due to time constraints, some features might not be tested thoroughly enough.',
         'Mitigation: Focus on testing the most critical user flows first. Use risk-based testing to prioritize high-risk areas.'),
        
        ('Scope Creep',
         'Risk: New features or requirements might be added during testing, expanding the scope.',
         'Mitigation: Stick to the original project scope. Document any new requirements separately and defer if possible.')
    ]
    
    for risk_title, risk_desc, mitigation in time_risks:
        p = doc.add_paragraph(f'{risk_title}: ', style='List Bullet')
        p.add_run(risk_desc)
        p = doc.add_paragraph(f'Mitigation: {mitigation}', style='List Bullet 2')
    
    add_heading_with_style(doc, '9.3 Possible Bugs and Issues', 2)
    doc.add_paragraph(
        'Based on the complexity of the project, the following types of bugs are likely to be encountered:'
    )
    
    bug_types = [
        ('Authentication Bugs',
         'Issues with JWT token expiration, session management, or MFA verification. These are critical and need immediate attention.'),
        
        ('API Integration Bugs',
         'Problems with external API calls, error handling when APIs fail, or incorrect data parsing from API responses.'),
        
        ('UI Bugs',
         'Layout issues on different screen sizes, form validation not working correctly, or buttons not responding.'),
        
        ('Data Display Bugs',
         'Incorrect data shown in reports, wrong statistics on dashboard, or missing information in scan results.'),
        
        ('Permission Bugs',
         'Users accessing features they should not have access to, or role-based restrictions not working correctly.')
    ]
    
    for bug_type, desc in bug_types:
        p = doc.add_paragraph(f'{bug_type}: ', style='List Bullet')
        p.add_run(desc)
    
    add_heading_with_style(doc, '9.4 Risk Handling Strategy', 2)
    doc.add_paragraph(
        'When risks materialize or bugs are found, the following strategy will be followed:'
    )
    
    strategy_steps = [
        'Document the issue clearly with steps to reproduce',
        'Assess the severity: Critical (blocks core functionality), High (major feature broken), Medium (minor issue), Low (cosmetic)',
        'Prioritize fixing critical and high-severity bugs first',
        'For non-critical bugs, decide whether to fix immediately or document for later',
        'Update test cases if requirements change or bugs reveal missing test scenarios',
        'Keep a bug log to track all issues and their resolution status'
    ]
    
    for step in strategy_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    # 10. Test Deliverables
    add_heading_with_style(doc, '10. Test Deliverables', 1)
    
    doc.add_paragraph(
        'The following documents and artifacts will be produced as part of the testing process:'
    )
    
    add_heading_with_style(doc, '10.1 Test Plan', 2)
    doc.add_paragraph(
        'This document (the Test Plan Report) serves as the master test plan. It outlines the testing strategy, '
        'test cases, schedule, and approach. It will be reviewed by the supervisor and used as a reference throughout testing.'
    )
    
    add_heading_with_style(doc, '10.2 Test Cases', 2)
    doc.add_paragraph(
        'A detailed list of test cases (like the sample test cases in section 6.3) will be maintained. Each test case '
        'will be executed, and results will be recorded (Pass/Fail/Blocked). Test cases may be stored in a spreadsheet '
        'or document for easy tracking.'
    )
    
    add_heading_with_style(doc, '10.3 Bug Reports', 2)
    doc.add_paragraph(
        'When bugs are found during testing, they will be documented in bug reports. Each bug report should include:'
    )
    
    bug_report_items = [
        'Bug ID (unique identifier)',
        'Title (brief description)',
        'Severity (Critical/High/Medium/Low)',
        'Steps to reproduce',
        'Expected behavior',
        'Actual behavior',
        'Screenshots or error messages (if applicable)',
        'Status (Open/Fixed/Verified/Deferred)',
        'Date found and date fixed'
    ]
    
    for item in bug_report_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_with_style(doc, '10.4 Final Test Summary', 2)
    doc.add_paragraph(
        'At the end of the testing phase, a test summary report will be prepared. This will include:'
    )
    
    summary_items = [
        'Total number of test cases executed',
        'Number of test cases passed, failed, and blocked',
        'Summary of bugs found (by severity)',
        'Overall assessment of system quality',
        'Recommendations for improvements',
        'List of known issues or limitations'
    ]
    
    for item in summary_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 11. Test Traceability Matrix
    add_heading_with_style(doc, '11. Test Traceability Matrix', 1)
    
    doc.add_paragraph(
        'A test traceability matrix helps ensure that all requirements are covered by test cases. This section maps '
        'high-level requirements to specific test cases.'
    )
    
    doc.add_paragraph(
        'Note: For a graduation project, requirements might be informal (based on features implemented). This matrix '
        'maps features/functionality to test cases to ensure nothing is missed.'
    )
    
    traceability_data = [
        ('REQ-001', 'User Authentication', 'TC-001, TC-002, TC-003, TC-004, TC-012'),
        ('REQ-002', 'Multi-Factor Authentication', 'TC-013'),
        ('REQ-003', 'Role-Based Access Control', 'TC-011'),
        ('REQ-004', 'Email Scan Creation', 'TC-005'),
        ('REQ-005', 'Domain Scan Creation', 'TC-006'),
        ('REQ-006', 'IP Address Scan Creation', 'TC-006 (similar to domain)'),
        ('REQ-007', 'View Scan Results', 'TC-007'),
        ('REQ-008', 'Generate PDF Reports', 'TC-008'),
        ('REQ-009', 'Generate HTML Reports', 'TC-008 (format variation)'),
        ('REQ-010', 'User Management (Admin)', 'TC-009, TC-010'),
        ('REQ-011', 'API Key Management (Admin)', 'TC-015'),
        ('REQ-012', 'Dashboard Statistics', 'TC-014'),
        ('REQ-013', 'Notification System', 'TC-016'),
        ('REQ-014', 'Responsive Design', 'TC-019'),
        ('REQ-015', 'Error Handling', 'TC-018, TC-020'),
        ('REQ-016', 'Input Validation', 'TC-004, TC-017'),
        ('REQ-017', 'Session Management', 'TC-020')
    ]
    
    headers = ['Requirement ID', 'Requirement/Feature', 'Related Test Cases']
    add_table_with_headers(doc, headers, traceability_data)
    
    doc.add_paragraph('')
    doc.add_paragraph(
        'This traceability matrix ensures that each major requirement has at least one test case covering it. '
        'During testing, if a requirement is not adequately covered, additional test cases can be added.'
    )
    
    # 12. Conclusion
    add_heading_with_style(doc, '12. Conclusion', 1)
    
    add_heading_with_style(doc, '12.1 Overall Testing Approach', 2)
    doc.add_paragraph(
        'The testing approach for CyberScope focuses on ensuring that all core functionality works correctly and that '
        'the system is secure and user-friendly. The approach is practical and tailored to a graduation project, '
        'balancing thoroughness with time constraints.'
    )
    doc.add_paragraph(
        'Testing will be done manually, with a focus on functional testing and basic security checks. The test cases '
        'cover both positive scenarios (things working correctly) and negative scenarios (error handling). Integration '
        'testing ensures that the frontend and backend work together, and that external API integrations function properly.'
    )
    doc.add_paragraph(
        'The testing process is iterative: as bugs are found and fixed, regression testing ensures that fixes do not '
        'break existing functionality. This approach helps maintain system stability throughout the development and testing phases.'
    )
    
    add_heading_with_style(doc, '12.2 Expected Outcomes', 2)
    doc.add_paragraph(
        'By the end of the testing phase, we expect to have:'
    )
    
    outcomes = [
        'All critical user flows tested and verified to work correctly',
        'Most bugs identified and fixed (or documented if deferred)',
        'A stable system that can be demonstrated to supervisors and evaluators',
        'Comprehensive documentation of test results and any known issues',
        'Confidence that the system meets the project requirements and is ready for evaluation'
    ]
    
    for outcome in outcomes:
        doc.add_paragraph(outcome, style='List Bullet')
    
    add_heading_with_style(doc, '12.3 Final Remarks', 2)
    doc.add_paragraph(
        'This test plan provides a structured approach to testing the CyberScope OSINT platform. While it may not be '
        'as comprehensive as a commercial software testing plan, it covers the essential aspects needed for a graduation project.'
    )
    doc.add_paragraph(
        'Testing is an important part of software development, and this plan helps ensure that the CyberScope platform '
        'is reliable, secure, and user-friendly. The test cases and strategies outlined here will guide the testing process '
        'and help identify areas that need improvement.'
    )
    doc.add_paragraph(
        'As testing progresses, this plan may be updated to reflect new findings or changes in requirements. Flexibility '
        'is important, but having a plan provides direction and helps ensure nothing important is overlooked.'
    )
    doc.add_paragraph(
        'The goal is to deliver a high-quality OSINT platform that demonstrates both technical skills and attention to '
        'detail in software testing and quality assurance.'
    )
    
    # Save document
    output_path = 'Test_Plan_Report_CyberScope.docx'
    doc.save(output_path)
    print(f'Test Plan Report saved to {output_path}')

if __name__ == '__main__':
    main()

